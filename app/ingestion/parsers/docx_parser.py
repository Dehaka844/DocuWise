from io import BytesIO

from docx import Document
from fastapi import UploadFile

from app.ingestion.parsers.base import BaseDocumentParser
from app.models.document import PageContent, ParsedDocument


class DOCXParser(BaseDocumentParser):

    async def parse(
        self,
        file: UploadFile,
    ) -> ParsedDocument:

        file_content = await file.read()

        document = Document(
            BytesIO(file_content)
        )

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        content = "\n".join(paragraphs)

        return ParsedDocument(
            pages=[
                PageContent(
                    page_number=1,
                    content=content,
                )
            ],
            metadata={
                "filename": file.filename,
                "source_type": "docx",
                "page_count": 1,
            },
        )